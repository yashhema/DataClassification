# content_extractor_validation_suite.py
"""
Comprehensive Content Extraction Validation Suite

This module validates actual content extraction accuracy, not just framework functionality.
Follows rigorous validation checklist to ensure content fidelity, encoding preservation,
and component classification accuracy.
"""

import asyncio
import os
import tempfile
import json
import zipfile
import tarfile
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Set, Tuple
import unicodedata

# Import your existing modules
from core.config.configuration_manager import ConfigurationManager
from core.db.database_interface import DatabaseInterface
from core.errors import ErrorHandler
from core.logging.system_logger import SystemLogger
from content_extraction.content_extractor import ContentExtractor

class TestDefinition:
    """Defines expected outcomes for a test file"""
    def __init__(self, file_path: str, description: str):
        self.file_path = file_path
        self.description = description
        self.expected_components = []
        self.expected_total_count = 0
        self.must_contain_exact = []  # Character-exact matches required
        self.must_contain_patterns = []  # Regex patterns that must match
        self.encoding_tests = []  # Specific encoding validation requirements
        self.structure_tests = []  # JSON/CSV/Table structure requirements
        self.anti_patterns = []  # Things that should NOT happen
        self.edge_case_behaviors = []  # Expected behavior for edge cases
    
    def add_expected_component(self, component_type: str, extraction_method: str, 
                             content_requirements: List[str], encoding: str = "utf-8"):
        """Add a component expectation with validation requirements"""
        self.expected_components.append({
            "type": component_type,
            "method": extraction_method,
            "content_requirements": content_requirements,
            "encoding": encoding
        })
        self.expected_total_count += 1
        return self
    
    def add_exact_content(self, content: str, encoding: str = "utf-8"):
        """Add content that must be extracted character-exactly"""
        self.must_contain_exact.append({"content": content, "encoding": encoding})
        return self
    
    def add_exact_content_simple(self, content: str):
        """Add content that must be extracted character-exactly (UTF-8 default)"""
        self.must_contain_exact.append({"content": content, "encoding": "utf-8"})
        return self
    
    def add_structure_test(self, structure_type: str, requirements: Dict):
        """Add structural validation requirements (tables, JSON, etc.)"""
        self.structure_tests.append({"type": structure_type, "requirements": requirements})
        return self
    
    def add_anti_pattern(self, pattern_type: str, description: str):
        """Add patterns that should NOT occur (duplications, misclassifications)"""
        self.anti_patterns.append({"type": pattern_type, "description": description})
        return self
    
    def add_encoding_test(self, original_chars: str, encoding: str):
        """Add specific encoding preservation test"""
        self.encoding_tests.append({"chars": original_chars, "encoding": encoding})
        return self

class ContentValidator:
    """Validates extracted components against expected outcomes"""
    
    def __init__(self):
        self.validation_results = {}
    
    def _validate_component_specific_content(self, components: List, test_def: TestDefinition, results: Dict):
        """Validate that each component contains its expected content"""
        results["component_content_analysis"] = {
            "component_matches": [],
            "validation_errors": [],
            "unmatched_expectations": [],
            "unexpected_components": []
        }
        
        # Group expected components by type+method to handle duplicates
        expected_groups = {}
        for i, expected in enumerate(test_def.expected_components):
            key = f"{expected['type']}#{expected['method']}"
            if key not in expected_groups:
                expected_groups[key] = []
            expected_groups[key].append((i, expected))
        
        # Group actual components by type+method
        actual_groups = {}
        for i, actual_comp in enumerate(components):
            actual_type = actual_comp.component_type
            actual_method = getattr(actual_comp, 'extraction_method', 'unknown')
            key = f"{actual_type}#{actual_method}"
            if key not in actual_groups:
                actual_groups[key] = []
            actual_groups[key].append((i, actual_comp))
        
        # Match groups and validate content
        for group_key, expected_list in expected_groups.items():
            actual_list = actual_groups.get(group_key, [])
            
            if len(expected_list) != len(actual_list):
                error = f"Component count mismatch for {group_key}: expected {len(expected_list)}, got {len(actual_list)}"
                results["component_content_analysis"]["validation_errors"].append(error)
                results["validation_errors"].append(error)
                results["validation_passed"] = False
                continue
            
            # For groups with multiple components, match by content similarity
            if len(expected_list) > 1:
                self._match_components_by_content(expected_list, actual_list, results)
            else:
                # Single component - direct match
                expected_idx, expected = expected_list[0]
                actual_idx, actual_comp = actual_list[0]
                self._validate_single_component_content(expected, actual_comp, results)
        
        # Check for unexpected component groups
        for group_key, actual_list in actual_groups.items():
            if group_key not in expected_groups:
                for actual_idx, actual_comp in actual_list:
                    results["component_content_analysis"]["unexpected_components"].append({
                        "component_id": actual_comp.component_id,
                        "type": actual_comp.component_type,
                        "method": getattr(actual_comp, 'extraction_method', 'unknown'),
                        "content_size": len(actual_comp.content) if actual_comp.content else 0
                    })
    
    def _match_components_by_content(self, expected_list: List, actual_list: List, results: Dict):
        """Match components with same type/method by content similarity"""
        used_actual = set()
        
        for expected_idx, expected in expected_list:
            best_match = None
            best_score = 0
            best_actual_idx = -1
            
            content_requirements = expected["content_requirements"]
            
            for actual_idx, actual_comp in actual_list:
                if actual_idx in used_actual:
                    continue
                
                actual_content = actual_comp.content or ""
                score = 0
                
                # Score based on how many content requirements are met
                for required_content in content_requirements:
                    if required_content in actual_content:
                        score += 1
                
                if score > best_score:
                    best_match = actual_comp
                    best_score = score
                    best_actual_idx = actual_idx
            
            if best_match:
                used_actual.add(best_actual_idx)
                self._validate_single_component_content(expected, best_match, results)
            else:
                results["component_content_analysis"]["unmatched_expectations"].append(expected)
                error = f"No suitable content match found for expected component {expected['type']}#{expected['method']}"
                results["component_content_analysis"]["validation_errors"].append(error)
                results["validation_errors"].append(error)
                results["validation_passed"] = False
    
    def _validate_single_component_content(self, expected: Dict, actual_comp, results: Dict):
        """Validate content requirements for a single component"""
        content_requirements = expected["content_requirements"]
        actual_content = actual_comp.content or ""
        
        component_validation = {
            "expected": expected,
            "actual_component_id": actual_comp.component_id,
            "content_requirements_met": [],
            "content_requirements_failed": [],
            "validation_passed": True
        }
        
        for required_content in content_requirements:
            if required_content in actual_content:
                component_validation["content_requirements_met"].append(required_content)
            else:
                component_validation["content_requirements_failed"].append(required_content)
                component_validation["validation_passed"] = False
                
                error = f"Component {actual_comp.component_id} missing expected content: '{required_content[:100]}...'"
                results["component_content_analysis"]["validation_errors"].append(error)
                results["validation_errors"].append(error)
                results["validation_passed"] = False
        
        results["component_content_analysis"]["component_matches"].append(component_validation)
    
    def validate_components(self, components: List, test_definition: TestDefinition) -> Dict[str, Any]:
        """Comprehensive validation of extracted components"""
        results = {
            "test_file": test_definition.file_path,
            "description": test_definition.description,
            "validation_passed": True,
            "validation_errors": [],
            "validation_warnings": [],
            "component_analysis": {},
            "content_analysis": {},
            "encoding_analysis": {},
            "structure_analysis": {},
            "anti_pattern_analysis": {}
        }
        
        # Component count validation
        self._validate_component_count(components, test_definition, results)
        
        # Component type and method validation
        self._validate_component_types(components, test_definition, results)
        
        # CRITICAL: Component-specific content validation
        self._validate_component_specific_content(components, test_definition, results)
        
        # Content accuracy validation (global requirements)
        self._validate_content_accuracy(components, test_definition, results)
        
        # Encoding preservation validation
        self._validate_encoding_preservation(components, test_definition, results)
        
        # Structure validation (tables, JSON, etc.)
        self._validate_structure_preservation(components, test_definition, results)
        
        # Anti-pattern detection
        self._validate_anti_patterns(components, test_definition, results)
        
        return results
    
    def _validate_component_count(self, components: List, test_def: TestDefinition, results: Dict):
        """Validate exact component count matches expectation"""
        expected_count = test_def.expected_total_count
        actual_count = len(components)
        
        results["component_analysis"]["expected_count"] = expected_count
        results["component_analysis"]["actual_count"] = actual_count
        results["component_analysis"]["count_match"] = (actual_count == expected_count)
        
        if actual_count != expected_count:
            error = f"Component count mismatch: expected {expected_count}, got {actual_count}"
            results["validation_errors"].append(error)
            results["validation_passed"] = False
    
    def _validate_component_types(self, components: List, test_def: TestDefinition, results: Dict):
        """Validate component types and extraction methods (order-independent)"""
        expected_components = test_def.expected_components
        expected_types = [comp["type"] for comp in expected_components]
        expected_methods = [comp["method"] for comp in expected_components]
        
        actual_types = [comp.component_type for comp in components]
        actual_methods = [getattr(comp, 'extraction_method', 'unknown') for comp in components]
        
        results["component_analysis"]["expected_types"] = expected_types
        results["component_analysis"]["actual_types"] = actual_types
        results["component_analysis"]["expected_methods"] = expected_methods
        results["component_analysis"]["actual_methods"] = actual_methods
        
        # Order-independent validation: check if we have the right types and methods
        expected_type_counts = {}
        expected_method_counts = {}
        for comp in expected_components:
            comp_type = comp["type"]
            comp_method = comp["method"]
            expected_type_counts[comp_type] = expected_type_counts.get(comp_type, 0) + 1
            expected_method_counts[comp_method] = expected_method_counts.get(comp_method, 0) + 1
        
        actual_type_counts = {}
        actual_method_counts = {}
        for comp in components:
            comp_type = comp.component_type
            comp_method = getattr(comp, 'extraction_method', 'unknown')
            actual_type_counts[comp_type] = actual_type_counts.get(comp_type, 0) + 1
            actual_method_counts[comp_method] = actual_method_counts.get(comp_method, 0) + 1
        
        results["component_analysis"]["expected_type_counts"] = expected_type_counts
        results["component_analysis"]["actual_type_counts"] = actual_type_counts
        results["component_analysis"]["expected_method_counts"] = expected_method_counts
        results["component_analysis"]["actual_method_counts"] = actual_method_counts
        
        # Check type count matches
        for expected_type, expected_count in expected_type_counts.items():
            actual_count = actual_type_counts.get(expected_type, 0)
            if actual_count != expected_count:
                error = f"Component type count mismatch for '{expected_type}': expected {expected_count}, got {actual_count}"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
        
        # Check for unexpected types
        for actual_type, actual_count in actual_type_counts.items():
            if actual_type not in expected_type_counts:
                warning = f"Unexpected component type '{actual_type}' found ({actual_count} instances)"
                results["validation_warnings"].append(warning)
        
        # Check method count matches  
        for expected_method, expected_count in expected_method_counts.items():
            actual_count = actual_method_counts.get(expected_method, 0)
            if actual_count != expected_count:
                error = f"Extraction method count mismatch for '{expected_method}': expected {expected_count}, got {actual_count}"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
    
    def _validate_content_accuracy(self, components: List, test_def: TestDefinition, results: Dict):
        """Validate content matches exactly with encoding awareness"""
        results["content_analysis"]["total_extracted_chars"] = 0
        results["content_analysis"]["exact_matches"] = []
        results["content_analysis"]["missing_content"] = []
        results["content_analysis"]["content_fidelity_tests"] = []
        
        # Build complete content map from all components
        component_contents = {}
        all_content = ""
        
        for comp in components:
            if comp.content:
                component_contents[comp.component_id] = comp.content
                all_content += comp.content + "\n"
                results["content_analysis"]["total_extracted_chars"] += len(comp.content)
        
        # Check exact content matches with character-level validation
        for exact_requirement in test_def.must_contain_exact:
            required_content = exact_requirement["content"]
            required_encoding = exact_requirement["encoding"]
            
            # Find which component(s) contain this content
            found_in_components = []
            content_found = False
            
            for comp_id, content in component_contents.items():
                if required_content in content:
                    found_in_components.append(comp_id)
                    content_found = True
                    
                    # For embedded content, just verify presence with proper context
                    # Don't require character-exact match for content that's part of larger text
                    break  # Found it, that's sufficient
            
            if found_in_components:
                # Verify encoding preservation
                try:
                    # Test that the content can be encoded/decoded properly
                    test_bytes = required_content.encode(required_encoding)
                    decoded_back = test_bytes.decode(required_encoding)
                    encoding_preserved = (decoded_back == required_content)
                    
                    results["content_analysis"]["exact_matches"].append({
                        "content": required_content,
                        "encoding": required_encoding,
                        "found_in_components": found_in_components,
                        "content_found": content_found,
                        "encoding_preserved": encoding_preserved,
                        "byte_length": len(test_bytes)
                    })
                    
                    if not encoding_preserved:
                        error = f"Encoding not preserved for content: '{required_content}' ({required_encoding})"
                        results["validation_errors"].append(error)
                        results["validation_passed"] = False
                        
                except UnicodeEncodeError as e:
                    error = f"Encoding error for content '{required_content}': {str(e)}"
                    results["validation_errors"].append(error)
                    results["validation_passed"] = False
            else:
                results["content_analysis"]["missing_content"].append(required_content)
                error = f"Required content not found: '{required_content}'"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
        
        # Additional content fidelity test: check for content corruption
        for comp_id, content in component_contents.items():
            fidelity_test = {
                "component_id": comp_id,
                "content_length": len(content),
                "has_null_bytes": '\x00' in content,
                "has_replacement_chars": '\ufffd' in content,
                "proper_line_endings": True
            }
            
            # Check for common corruption indicators
            if fidelity_test["has_null_bytes"]:
                error = f"Component {comp_id} contains null bytes (possible binary corruption)"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
            
            if fidelity_test["has_replacement_chars"]:
                error = f"Component {comp_id} contains Unicode replacement characters (encoding corruption)"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
            
            results["content_analysis"]["content_fidelity_tests"].append(fidelity_test)
    
    def _validate_encoding_preservation(self, components: List, test_def: TestDefinition, results: Dict):
        """Validate Unicode and encoding preservation"""
        results["encoding_analysis"]["tests_performed"] = []
        
        for encoding_test in test_def.encoding_tests:
            test_chars = encoding_test["chars"]
            test_encoding = encoding_test["encoding"]
            
            # Find these characters in extracted content
            found_chars = []
            for comp in components:
                if comp.content and test_chars in comp.content:
                    # Extract the specific characters and validate encoding
                    start_idx = comp.content.find(test_chars)
                    if start_idx != -1:
                        extracted_chars = comp.content[start_idx:start_idx + len(test_chars)]
                        found_chars.append(extracted_chars)
                        
                        # Byte-level comparison
                        try:
                            original_bytes = test_chars.encode(test_encoding)
                            extracted_bytes = extracted_chars.encode('utf-8')  # ContentExtractor should normalize to UTF-8
                            
                            # For UTF-8, compare both raw and normalized forms
                            if test_encoding.lower().startswith('utf'):
                                # First check raw byte-level comparison
                                raw_match = (test_chars == extracted_chars)
                                
                                # Then check normalized comparison
                                normalized_original = unicodedata.normalize('NFC', test_chars)
                                normalized_extracted = unicodedata.normalize('NFC', extracted_chars)
                                normalized_match = (normalized_original == normalized_extracted)
                                
                                # Encoding is preserved if either raw match OR normalized match with no corruption indicators
                                has_corruption = '\ufffd' in extracted_chars or '\x00' in extracted_chars
                                encoding_match = raw_match or (normalized_match and not has_corruption)
                                
                                if not raw_match and normalized_match:
                                    # Log normalization case for debugging
                                    results["encoding_analysis"]["normalization_applied"] = results["encoding_analysis"].get("normalization_applied", [])
                                    results["encoding_analysis"]["normalization_applied"].append({
                                        "test_chars": test_chars,
                                        "extracted_chars": extracted_chars,
                                        "raw_match": raw_match,
                                        "normalized_match": normalized_match
                                    })
                            else:
                                # For other encodings, compare after conversion
                                try:
                                    converted_original = test_chars.encode(test_encoding).decode(test_encoding)
                                    encoding_match = (converted_original == extracted_chars)
                                except:
                                    encoding_match = False
                            
                            results["encoding_analysis"]["tests_performed"].append({
                                "test_chars": test_chars,
                                "original_encoding": test_encoding,
                                "extracted_chars": extracted_chars,
                                "encoding_preserved": encoding_match,
                                "original_byte_length": len(original_bytes),
                                "extracted_byte_length": len(extracted_bytes)
                            })
                            
                            if not encoding_match:
                                error = f"Encoding not preserved for '{test_chars}': original '{test_chars}' != extracted '{extracted_chars}'"
                                results["validation_errors"].append(error)
                                results["validation_passed"] = False
                                
                        except Exception as e:
                            error = f"Encoding validation failed for '{test_chars}': {str(e)}"
                            results["validation_errors"].append(error)
                            results["validation_passed"] = False
            
            if not found_chars:
                error = f"Encoding test characters '{test_chars}' not found in any extracted content"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
    
    def _validate_structure_preservation(self, components: List, test_def: TestDefinition, results: Dict):
        """Validate structural data preservation (tables, JSON hierarchies)"""
        results["structure_analysis"]["tests_performed"] = []
        
        for structure_test in test_def.structure_tests:
            test_type = structure_test["type"]
            requirements = structure_test["requirements"]
            
            if test_type == "table":
                self._validate_table_structure(components, requirements, results)
            elif test_type == "json":
                self._validate_json_structure(components, requirements, results)
            elif test_type == "csv":
                # CSV can be processed as either table OR text component
                self._validate_csv_structure_flexible(components, requirements, results)
    
    def _validate_table_structure(self, components: List, requirements: Dict, results: Dict):
        """Validate table component structure"""
        table_components = [comp for comp in components if comp.component_type == "table"]
        
        if not table_components and requirements.get("required", True):
            error = "Expected table component but none found"
            results["validation_errors"].append(error)
            results["validation_passed"] = False
            return
        
        for table_comp in table_components:
            if not hasattr(table_comp, 'schema') or not table_comp.schema:
                error = f"Table component {table_comp.component_id} missing schema"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
                continue
            
            schema = table_comp.schema
            
            # Validate table dimensions
            if "expected_rows" in requirements:
                actual_rows = len(schema.get("rows", []))
                expected_rows = requirements["expected_rows"]
                if actual_rows != expected_rows:
                    error = f"Table row count mismatch: expected {expected_rows}, got {actual_rows}"
                    results["validation_errors"].append(error)
                    results["validation_passed"] = False
            
            if "expected_columns" in requirements:
                actual_columns = len(schema.get("headers", []))
                expected_columns = requirements["expected_columns"]
                if actual_columns != expected_columns:
                    error = f"Table column count mismatch: expected {expected_columns}, got {actual_columns}"
                    results["validation_errors"].append(error)
                    results["validation_passed"] = False
            
            # Validate specific cell content
            if "expected_cells" in requirements:
                rows = schema.get("rows", [])
                for cell_requirement in requirements["expected_cells"]:
                    row_idx = cell_requirement["row"]
                    col_name = cell_requirement["column"]
                    expected_value = cell_requirement["value"]
                    
                    if row_idx < len(rows) and col_name in rows[row_idx]:
                        actual_value = rows[row_idx][col_name]
                        if actual_value != expected_value:
                            error = f"Table cell mismatch at row {row_idx}, column '{col_name}': expected '{expected_value}', got '{actual_value}'"
                            results["validation_errors"].append(error)
                            results["validation_passed"] = False
                    else:
                        error = f"Table cell not found: row {row_idx}, column '{col_name}'"
                        results["validation_errors"].append(error)
                        results["validation_passed"] = False
            
            results["structure_analysis"]["tests_performed"].append({
                "type": "table",
                "component_id": table_comp.component_id,
                "rows_found": len(schema.get("rows", [])),
                "columns_found": len(schema.get("headers", [])),
                "headers": schema.get("headers", [])
            })
    
    def _validate_json_structure(self, components: List, requirements: Dict, results: Dict):
        """Validate JSON structure preservation"""
        # For JSON files processed as text, try to parse the content
        for comp in components:
            if comp.content and requirements.get("should_be_parseable", True):
                try:
                    parsed_data = json.loads(comp.content)
                    
                    # Validate specific JSON structure requirements
                    if "expected_keys" in requirements:
                        for key_path in requirements["expected_keys"]:
                            if not self._check_json_key_path(parsed_data, key_path):
                                error = f"JSON key path '{key_path}' not found in parsed content"
                                results["validation_errors"].append(error)
                                results["validation_passed"] = False
                    
                    results["structure_analysis"]["tests_performed"].append({
                        "type": "json",
                        "component_id": comp.component_id,
                        "successfully_parsed": True,
                        "top_level_keys": list(parsed_data.keys()) if isinstance(parsed_data, dict) else []
                    })
                    
                except json.JSONDecodeError as e:
                    error = f"JSON content not parseable: {str(e)}"
                    results["validation_errors"].append(error)
                    results["validation_passed"] = False
    
    def _check_json_key_path(self, data: Any, key_path: str) -> bool:
        """Check if a dot-separated key path exists in JSON data"""
        keys = key_path.split('.')
        current = data
        
        for key in keys:
            if isinstance(current, dict) and key in current:
                current = current[key]
            else:
                return False
        return True
    
    def _validate_csv_structure_flexible(self, components: List, requirements: Dict, results: Dict):
        """Validate CSV structure whether processed as table or text"""
        csv_validated = False
        
        # First try table components
        table_components = [comp for comp in components if comp.component_type == "table"]
        if table_components:
            for table_comp in table_components:
                if hasattr(table_comp, 'schema') and table_comp.schema:
                    self._validate_table_structure([table_comp], requirements, results)
                    csv_validated = True
                    break
        
        # If no table found, try text components that look like CSV
        if not csv_validated:
            for comp in components:
                if comp.content and "\n" in comp.content and "," in comp.content:
                    lines = comp.content.strip().split('\n')
                    if len(lines) >= 2:  # Header + at least one data row
                        header_line = lines[0]
                        data_lines = [line for line in lines[1:] if line.strip()]
                        
                        header_count = len(header_line.split(','))
                        
                        if "expected_columns" in requirements:
                            expected_cols = requirements["expected_columns"]
                            if header_count != expected_cols:
                                error = f"CSV column count mismatch: expected {expected_cols}, got {header_count}"
                                results["validation_errors"].append(error)
                                results["validation_passed"] = False
                            else:
                                results["structure_analysis"]["tests_performed"].append({
                                    "type": "csv_as_text",
                                    "component_id": comp.component_id,
                                    "columns_detected": header_count,
                                    "data_rows_detected": len(data_lines),
                                    "validation_passed": True
                                })
                        
                        if "expected_rows" in requirements:
                            expected_rows = requirements["expected_rows"]
                            if len(data_lines) != expected_rows:
                                error = f"CSV row count mismatch: expected {expected_rows}, got {len(data_lines)}"
                                results["validation_errors"].append(error)
                                results["validation_passed"] = False
                        
                        csv_validated = True
                        break
        
        if not csv_validated:
            error = "CSV structure validation failed: no table or recognizable CSV text found"
            results["validation_errors"].append(error)
            results["validation_passed"] = False
    
    def _validate_csv_structure(self, components: List, requirements: Dict, results: Dict):
        """Validate CSV structure preservation"""
        # CSV might be processed as text or table
        for comp in components:
            if comp.component_type == "table":
                # Already validated by table structure validation
                continue
            elif comp.content and "\n" in comp.content and "," in comp.content:
                # Likely CSV processed as text - validate structure
                lines = comp.content.strip().split('\n')
                if len(lines) >= 2:  # Header + at least one data row
                    header_line = lines[0]
                    data_lines = lines[1:]
                    
                    header_count = len(header_line.split(','))
                    
                    if "expected_columns" in requirements:
                        expected_cols = requirements["expected_columns"]
                        if header_count != expected_cols:
                            error = f"CSV column count mismatch: expected {expected_cols}, got {header_count}"
                            results["validation_errors"].append(error)
                            results["validation_passed"] = False
                    
                    # Check data row consistency
                    for i, line in enumerate(data_lines):
                        if line.strip():  # Skip empty lines
                            data_count = len(line.split(','))
                            if data_count != header_count:
                                warning = f"CSV row {i+1} has {data_count} fields, header has {header_count}"
                                results["validation_warnings"].append(warning)
                    
                    results["structure_analysis"]["tests_performed"].append({
                        "type": "csv",
                        "component_id": comp.component_id,
                        "columns_detected": header_count,
                        "data_rows_detected": len([l for l in data_lines if l.strip()])
                    })
    
    def _validate_anti_patterns(self, components: List, test_def: TestDefinition, results: Dict):
        """Detect patterns that should not occur"""
        results["anti_pattern_analysis"]["detected_issues"] = []
        
        for anti_pattern in test_def.anti_patterns:
            pattern_type = anti_pattern["type"]
            
            if pattern_type == "content_duplication":
                self._detect_content_duplication(components, results)
            elif pattern_type == "type_misclassification":
                self._detect_type_misclassification(components, anti_pattern, results)
            elif pattern_type == "archive_incompleteness":
                self._detect_archive_incompleteness(components, anti_pattern, results)
    
    def _detect_content_duplication(self, components: List, results: Dict):
        """Detect identical or overlapping content between components"""
        content_hashes = {}
        duplications = []
        
        for i, comp in enumerate(components):
            if not comp.content or len(comp.content.strip()) < 20:
                continue
            
            # Create hash of significant content chunks
            content_hash = hashlib.md5(comp.content.encode()).hexdigest()
            
            if content_hash in content_hashes:
                duplications.append({
                    "type": "exact_duplicate",
                    "component1": content_hashes[content_hash],
                    "component2": comp.component_id,
                    "content_preview": comp.content[:100]
                })
            else:
                content_hashes[content_hash] = comp.component_id
            
            # Check for substantial overlap with other components
            words = comp.content.split()
            if len(words) > 15:
                for j, other_comp in enumerate(components):
                    if i != j and other_comp.content and len(other_comp.content.split()) > 15:
                        overlap_ratio = self._calculate_content_overlap(comp.content, other_comp.content)
                        if overlap_ratio > 0.7:  # 70% overlap threshold
                            duplications.append({
                                "type": "substantial_overlap",
                                "component1": comp.component_id,
                                "component2": other_comp.component_id,
                                "overlap_ratio": overlap_ratio
                            })
        
        if duplications:
            results["anti_pattern_analysis"]["detected_issues"].extend(duplications)
            for dup in duplications:
                error = f"Content duplication detected: {dup['component1']} ↔ {dup['component2']} ({dup['type']})"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
    
    def _calculate_content_overlap(self, content1: str, content2: str) -> float:
        """Calculate overlap ratio between two content strings"""
        words1 = set(content1.lower().split())
        words2 = set(content2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0
    
    def _detect_type_misclassification(self, components: List, anti_pattern: Dict, results: Dict):
        """Detect component type misclassifications"""
        expected_types = anti_pattern.get("expected_types", [])
        
        for comp in components:
            actual_type = comp.component_type
            
            # Check if JSON/CSV files are being processed as plain text
            if "json" in os.path.basename(comp.parent_path).lower() and actual_type == "text":
                if getattr(comp, 'extraction_method', '') == "text_file_read":
                    issue = {
                        "type": "json_as_text",
                        "component_id": comp.component_id,
                        "file_path": comp.parent_path,
                        "actual_type": actual_type,
                        "expected_type": "json_parsed"
                    }
                    results["anti_pattern_analysis"]["detected_issues"].append(issue)
                    warning = f"JSON file processed as plain text: {comp.component_id}"
                    results["validation_warnings"].append(warning)
            
            elif "csv" in os.path.basename(comp.parent_path).lower() and actual_type == "text":
                if getattr(comp, 'extraction_method', '') == "text_file_read":
                    issue = {
                        "type": "csv_as_text",
                        "component_id": comp.component_id,
                        "file_path": comp.parent_path,
                        "actual_type": actual_type,
                        "expected_type": "table"
                    }
                    results["anti_pattern_analysis"]["detected_issues"].append(issue)
                    warning = f"CSV file processed as plain text instead of table: {comp.component_id}"
                    results["validation_warnings"].append(warning)
    
    def _detect_archive_incompleteness(self, components: List, anti_pattern: Dict, results: Dict):
        """Detect incomplete archive processing"""
        archive_components = [comp for comp in components if comp.component_type == "archive_member"]
        
        if archive_components:
            # Check if archive processing hit limits
            expected_member_count = anti_pattern.get("expected_member_count")
            actual_member_count = len(archive_components)
            
            if expected_member_count and actual_member_count < expected_member_count:
                issue = {
                    "type": "archive_incomplete",
                    "expected_members": expected_member_count,
                    "actual_members": actual_member_count,
                    "difference": expected_member_count - actual_member_count
                }
                results["anti_pattern_analysis"]["detected_issues"].append(issue)
                error = f"Archive processing incomplete: {actual_member_count}/{expected_member_count} members processed"
                results["validation_errors"].append(error)
                results["validation_passed"] = False
            
            # Additional validation: check if we got the expected member names
            if "expected_member_names" in anti_pattern:
                expected_names = set(anti_pattern["expected_member_names"])
                actual_names = set()
                
                for comp in archive_components:
                    member_name = comp.metadata.get("member_name", "unknown")
                    if member_name != "unknown":
                        actual_names.add(member_name)
                
                missing_members = expected_names - actual_names
                if missing_members:
                    error = f"Archive missing expected members: {missing_members}"
                    results["validation_errors"].append(error)
                    results["validation_passed"] = False

class TestFileCreator:
    """Creates test files with known content for validation"""
    
    def __init__(self, test_dir: str):
        self.test_dir = test_dir
        self.test_definitions = []
        
        # Standard test strings with various encodings
        self.unicode_test_strings = {
            "basic_latin": "Hello World",
            "accented_latin": "Café résumé naïve",
            "extended_latin": "Zürich Köln Ñoño",
            "currency_symbols": "€100 £50 ¥1000",
            "mathematical": "α² + β² = γ²",
            "emoji": "🚀 📱 💻 🔒",
            "cjk": "你好世界 こんにちは 안녕하세요",
            "mixed": "Café 🚀 你好 €100"
        }
        
        # DLP test data patterns
        self.dlp_patterns = {
            "ssn": "123-45-6789",
            "credit_card": "4532-1234-5678-9012",
            "email": "test@example.com",
            "phone": "555-123-4567"
        }
    
    def _verify_test_file_creation(self, file_path: str, expected_content: str, encoding: str = "utf-8") -> bool:
        """Verify that a created test file contains the expected content"""
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                actual_content = f.read()
            
            if actual_content != expected_content:
                print(f"WARNING: Test file creation verification failed for {file_path}")
                print(f"Expected: {expected_content[:100]}...")
                print(f"Actual: {actual_content[:100]}...")
                return False
            return True
        except Exception as e:
            print(f"ERROR: Could not verify test file {file_path}: {e}")
            return False
    
    def create_encoding_test_files(self) -> List[TestDefinition]:
        """Create files with different encodings to test preservation"""
        encoding_tests = []
        
        # Test 1: UTF-8 without BOM
        test_content = f"UTF-8 Test: {self.unicode_test_strings['mixed']} SSN: {self.dlp_patterns['ssn']}"
        utf8_file = os.path.join(self.test_dir, "test_001_utf8.txt")
        with open(utf8_file, 'w', encoding='utf-8') as f:
            f.write(test_content)
        
        # ✅ CHECKLIST: Test File Creation Validation
        if not self._verify_test_file_creation(utf8_file, test_content, 'utf-8'):
            raise ValueError(f"Test file creation failed for {utf8_file}")
        
        test_def = TestDefinition(utf8_file, "UTF-8 encoding preservation test")
        test_def.add_expected_component("text", "text_file_read", [test_content])
        test_def.add_exact_content(test_content, "utf-8")
        test_def.add_encoding_test(self.unicode_test_strings['mixed'], "utf-8")
        test_def.add_encoding_test(self.dlp_patterns['ssn'], "utf-8")
        encoding_tests.append(test_def)
        
        # Test 2: UTF-8 with BOM
        utf8_bom_file = os.path.join(self.test_dir, "test_002_utf8_bom.txt")
        with open(utf8_bom_file, 'w', encoding='utf-8-sig') as f:
            f.write(test_content)
        
        test_def = TestDefinition(utf8_bom_file, "UTF-8 with BOM encoding test")
        test_def.add_expected_component("text", "text_file_read", [test_content])
        test_def.add_exact_content(test_content, "utf-8")
        test_def.add_encoding_test(self.unicode_test_strings['mixed'], "utf-8")
        encoding_tests.append(test_def)
        
        # Test 3: UTF-16
        utf16_file = os.path.join(self.test_dir, "test_003_utf16.txt")
        with open(utf16_file, 'w', encoding='utf-16') as f:
            f.write(test_content)
        
        test_def = TestDefinition(utf16_file, "UTF-16 encoding test")
        test_def.add_expected_component("text", "text_file_read", [test_content])
        test_def.add_exact_content(test_content, "utf-16")
        test_def.add_encoding_test(self.unicode_test_strings['mixed'], "utf-16")
        encoding_tests.append(test_def)
        
        # Test 4: Windows-1252 (limited character set)
        win1252_content = f"Windows-1252 Test: {self.unicode_test_strings['accented_latin']} SSN: {self.dlp_patterns['ssn']}"
        win1252_file = os.path.join(self.test_dir, "test_004_win1252.txt")
        with open(win1252_file, 'w', encoding='windows-1252') as f:
            f.write(win1252_content)
        
        test_def = TestDefinition(win1252_file, "Windows-1252 encoding test")
        test_def.add_expected_component("text", "text_file_read", [win1252_content])
        test_def.add_exact_content(win1252_content, "windows-1252")
        test_def.add_encoding_test(self.unicode_test_strings['accented_latin'], "windows-1252")
        encoding_tests.append(test_def)
        
        self.test_definitions.extend(encoding_tests)
        return encoding_tests
    
    def create_structured_format_tests(self) -> List[TestDefinition]:
        """Create JSON, CSV, XML files to test structure preservation"""
        format_tests = []
        
        # Test 5: JSON structure test
        json_data = {
            "customer": {
                "name": f"{self.unicode_test_strings['accented_latin']}",
                "ssn": self.dlp_patterns['ssn'],
                "contact": {
                    "email": self.dlp_patterns['email'],
                    "phone": self.dlp_patterns['phone']
                }
            },
            "metadata": {
                "test_id": "JSON_STRUCTURE_005",
                "unicode_test": self.unicode_test_strings['emoji']
            }
        }
        
        json_file = os.path.join(self.test_dir, "test_005_structure.json")
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        test_def = TestDefinition(json_file, "JSON structure preservation test")
        test_def.add_expected_component("text", "json_parsing", [json.dumps(json_data, indent=2, ensure_ascii=False)])
        test_def.add_exact_content_simple(self.dlp_patterns['ssn'])
        test_def.add_exact_content_simple(self.unicode_test_strings['emoji'])
        test_def.add_structure_test("json", {
            "should_be_parseable": True,
            "expected_keys": ["customer.ssn", "customer.contact.email", "metadata.test_id"]
        })
        test_def.add_anti_pattern("type_misclassification", "JSON should not be processed as plain text")
        format_tests.append(test_def)
        
        # Test 6: CSV structure test
        csv_content = f"""Name,SSN,Email,Unicode_Test
{self.unicode_test_strings['accented_latin']},{self.dlp_patterns['ssn']},{self.dlp_patterns['email']},{self.unicode_test_strings['cjk']}
John Doe,987-65-4321,john@test.com,{self.unicode_test_strings['emoji']}
Jane Smith,111-22-3333,jane@test.com,CSV_TEST_006"""
        
        csv_file = os.path.join(self.test_dir, "test_006_structure.csv")
        with open(csv_file, 'w', encoding='utf-8') as f:
            f.write(csv_content)
        
        test_def = TestDefinition(csv_file, "CSV structure preservation test")
        test_def.add_expected_component("text", "text_file_read", [csv_content])  # May be processed as text
        test_def.add_exact_content_simple(self.dlp_patterns['ssn'])
        test_def.add_exact_content_simple(self.unicode_test_strings['cjk'])
        test_def.add_structure_test("csv", {
            "expected_columns": 4,
            "expected_rows": 3
        })
        test_def.add_anti_pattern("type_misclassification", "CSV should ideally be processed as table")
        format_tests.append(test_def)
        
        # Test 7: XML structure test
        xml_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<document>
    <metadata id="XML_TEST_007">
        <unicode_test>{self.unicode_test_strings['mixed']}</unicode_test>
    </metadata>
    <customer>
        <name>{self.unicode_test_strings['accented_latin']}</name>
        <ssn>{self.dlp_patterns['ssn']}</ssn>
        <contact>
            <email>{self.dlp_patterns['email']}</email>
        </contact>
    </customer>
</document>"""
        
        xml_file = os.path.join(self.test_dir, "test_007_structure.xml")
        with open(xml_file, 'w', encoding='utf-8') as f:
            f.write(xml_content)
        
        test_def = TestDefinition(xml_file, "XML structure preservation test")
        test_def.add_expected_component("text", "xml_parsing", [xml_content])
        test_def.add_exact_content_simple(self.dlp_patterns['ssn'])
        test_def.add_exact_content_simple(self.unicode_test_strings['mixed'])
        test_def.add_encoding_test(self.unicode_test_strings['mixed'], "utf-8")
        format_tests.append(test_def)
        
        self.test_definitions.extend(format_tests)
        return format_tests
    
    def create_office_document_tests(self) -> List[TestDefinition]:
        """Create office documents to test complex extraction"""
        office_tests = []
        
        # Test 8: Word document with table - duplication test
        try:
            from docx import Document
            docx_file = os.path.join(self.test_dir, "test_008_word_duplicate.docx")
            doc = Document()
            
            # Add paragraph with SSN
            doc.add_heading(f'Document Test - DOCX_008 {self.unicode_test_strings["emoji"]}', 0)
            para_text = f'Employee information: SSN {self.dlp_patterns["ssn"]} Email: {self.dlp_patterns["email"]}'
            doc.add_paragraph(para_text)
            
            # Add table with SAME SSN (test duplication detection)
            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).text = 'Name'
            table.cell(0, 1).text = 'SSN'
            table.cell(0, 2).text = 'Email'
            table.cell(1, 0).text = self.unicode_test_strings['accented_latin']
            table.cell(1, 1).text = self.dlp_patterns['ssn']  # SAME SSN as paragraph
            table.cell(1, 2).text = self.dlp_patterns['email']
            
            doc.save(docx_file)
            
            test_def = TestDefinition(docx_file, "Word document duplication detection test")
            test_def.add_expected_component("text", "python_docx_text", [para_text])
            test_def.add_expected_component("table", "python_docx_tables", [])
            test_def.add_exact_content(self.dlp_patterns['ssn'])
            test_def.add_exact_content(self.unicode_test_strings['emoji'])
            test_def.add_structure_test("table", {
                "expected_rows": 1,  # Data rows (excluding header)
                "expected_columns": 3,
                "expected_cells": [
                    {"row": 0, "column": "SSN", "value": self.dlp_patterns['ssn']}
                ]
            })
            test_def.add_anti_pattern("content_duplication", "Same SSN should not appear in both text and table components")
            office_tests.append(test_def)
            
        except ImportError:
            print("Warning: python-docx not available, skipping Word document test")
        
        # Test 9: Excel with multiple sheets
        try:
            import pandas as pd
            xlsx_file = os.path.join(self.test_dir, "test_009_excel_multisheet.xlsx")
            
            df1 = pd.DataFrame({
                'Name': [self.unicode_test_strings['accented_latin'], 'John Doe'],
                'SSN': [self.dlp_patterns['ssn'], '987-65-4321'],
                'Unicode': [self.unicode_test_strings['cjk'], self.unicode_test_strings['emoji']]
            })
            
            df2 = pd.DataFrame({
                'Credit_Card': [self.dlp_patterns['credit_card']],
                'Test_ID': ['EXCEL_009'],
                'Symbols': [self.unicode_test_strings['currency_symbols']]
            })
            
            with pd.ExcelWriter(xlsx_file) as writer:
                df1.to_excel(writer, sheet_name='Employees', index=False)
                df2.to_excel(writer, sheet_name='Financial', index=False)
            
            test_def = TestDefinition(xlsx_file, "Excel multi-sheet test")
            test_def.add_expected_component("text", "excel_cell_text", [])
            test_def.add_expected_component("table", "pandas_excel", [])
            test_def.add_expected_component("table", "pandas_excel", [])
            test_def.add_exact_content(self.dlp_patterns['ssn'])
            test_def.add_exact_content(self.dlp_patterns['credit_card'])
            test_def.add_exact_content(self.unicode_test_strings['cjk'])
            test_def.add_encoding_test(self.unicode_test_strings['currency_symbols'], "utf-8")
            office_tests.append(test_def)
            
        except ImportError:
            print("Warning: pandas not available, skipping Excel test")
        
        self.test_definitions.extend(office_tests)
        return office_tests
    
    def _validate_archive_completeness(self, file_path: str, expected_member_count: int) -> bool:
        """Verify archive actually contains expected number of members"""
        try:
            if file_path.endswith('.zip'):
                import zipfile
                with zipfile.ZipFile(file_path, 'r') as zf:
                    actual_count = len([name for name in zf.namelist() if not name.endswith('/')])
                    return actual_count == expected_member_count
            elif file_path.endswith('.tar'):
                import tarfile
                with tarfile.open(file_path, 'r') as tf:
                    actual_count = len([member for member in tf.getmembers() if member.isfile()])
                    return actual_count == expected_member_count
            return False
        except Exception as e:
            print(f"WARNING: Could not verify archive {file_path}: {e}")
            return False
    
    def create_archive_tests(self) -> List[TestDefinition]:
        """Create archive files to test completeness and classification"""
        archive_tests = []
        
        # Test 10: Simple ZIP with known members
        zip_file = os.path.join(self.test_dir, "test_010_simple_archive.zip")
        with zipfile.ZipFile(zip_file, 'w') as zf:
            # File 1: Text with Unicode
            text_content = f"Archive Text: {self.unicode_test_strings['mixed']} SSN: {self.dlp_patterns['ssn']}"
            zf.writestr("member1.txt", text_content)
            
            # File 2: JSON
            json_content = json.dumps({
                "archive_test": "ZIP_010",
                "ssn": self.dlp_patterns['ssn'],
                "unicode": self.unicode_test_strings['emoji']
            }, ensure_ascii=False)
            zf.writestr("member2.json", json_content)
            
            # File 3: CSV
            csv_content = f"Name,SSN\n{self.unicode_test_strings['accented_latin']},{self.dlp_patterns['ssn']}"
            zf.writestr("member3.csv", csv_content)
        
        test_def = TestDefinition(zip_file, "Simple ZIP archive test")
        test_def.add_expected_component("archive_member", "text_file_read", [text_content])
        test_def.add_expected_component("archive_member", "text_file_read", [json_content])
        test_def.add_expected_component("archive_member", "text_file_read", [csv_content])
        test_def.add_exact_content_simple(self.dlp_patterns['ssn'])
        test_def.add_exact_content_simple(self.unicode_test_strings['mixed'])
        test_def.add_anti_pattern("archive_incompleteness", {
            "expected_member_count": 3,
            "expected_member_names": ["member1.txt", "member2.json", "member3.csv"]
        })
        archive_tests.append(test_def)
        
        # Test 11: Large archive to test limits
        large_zip_file = os.path.join(self.test_dir, "test_011_large_archive.zip")
        with zipfile.ZipFile(large_zip_file, 'w') as zf:
            for i in range(75):  # Create 75 files to test 50-member limit
                content = f"File {i}: SSN {self.dlp_patterns['ssn']} Unicode: {self.unicode_test_strings['emoji']} ID: LARGE_011"
                zf.writestr(f"file_{i:03d}.txt", content)
        
        test_def = TestDefinition(large_zip_file, "Large archive limit test")
        # Expect only 50 components due to limit
        for i in range(50):
            test_def.add_expected_component("archive_member", "text_file_read", [])
        test_def.add_exact_content("LARGE_011")
        test_def.edge_case_behaviors.append({
            "type": "member_limit_enforcement",
            "expected_behavior": "Processing stops at 50 members"
        })
        archive_tests.append(test_def)
        
        self.test_definitions.extend(archive_tests)
        return archive_tests
    
    def create_ocr_tests(self) -> List[TestDefinition]:
        """Create image files for OCR testing"""
        ocr_tests = []
        
        # Test 12: Image with clear text for OCR
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            image_file = os.path.join(self.test_dir, "test_012_ocr_image.png")
            
            # Create high-contrast image with clear text
            img = Image.new('RGB', (600, 300), color='white')
            draw = ImageDraw.Draw(img)
            
            # Text to render in image
            ocr_text = f"OCR Test: SSN {self.dlp_patterns['ssn']}\nCredit Card: {self.dlp_patterns['credit_card']}\nTest ID: OCR_012"
            
            try:
                font = ImageFont.load_default()
            except:
                font = None
            
            # Draw text in large, clear font
            draw.text((50, 100), ocr_text, fill='black', font=font)
            img.save(image_file)
            
            test_def = TestDefinition(image_file, "OCR text extraction test")
            test_def.add_expected_component("image_ocr", "easyocr", [ocr_text])
            test_def.add_exact_content(self.dlp_patterns['ssn'])
            test_def.add_exact_content(self.dlp_patterns['credit_card'])
            test_def.add_exact_content("OCR_012")
            
            # Note: OCR may not be character-exact due to recognition limitations
            test_def.edge_case_behaviors.append({
                "type": "ocr_accuracy",
                "expected_behavior": "OCR text should contain key patterns even if not character-exact"
            })
            
            ocr_tests.append(test_def)
            
        except ImportError:
            print("Warning: PIL not available, skipping OCR test")
        
        self.test_definitions.extend(ocr_tests)
        return ocr_tests
    
    def get_all_test_definitions(self) -> List[TestDefinition]:
        """Return all created test definitions"""
        return self.test_definitions

class ContentExtractionValidator:
    """Main test orchestrator and validator"""
    
    def __init__(self, config_file_path: str = "config/system_default.yaml"):
        self.config_file_path = config_file_path
        self.validator = ContentValidator()
        self.test_results = []
    
    async def initialize_extractor(self):
        """Initialize ContentExtractor with proper configuration"""
        import logging
        logging.basicConfig(level="INFO", format='%(asctime)s - %(message)s')
        
        error_handler = ErrorHandler()
        temp_logger = SystemLogger(logging.getLogger("test_startup"), "TEXT", {})
        
        config_manager = ConfigurationManager(self.config_file_path)
        config_manager.set_core_services(temp_logger, error_handler)
        
        db_interface = DatabaseInterface(config_manager.get_db_connection_string(), temp_logger, error_handler)
        await db_interface.test_connection()
        
        await config_manager.load_database_overrides_async(db_interface)
        config_manager.finalize()
        settings = config_manager.settings
        
        logger = SystemLogger(logging.getLogger(__name__), settings.logging.format, config_manager.get_partial_ambient_context())
        db_interface.logger = logger
        
        self.content_extractor = ContentExtractor(
            extraction_config=settings.content_extraction,
            system_config=settings,
            logger=logger,
            error_handler=error_handler
        )
        
        await db_interface.close_async()
        return True
    
    async def validate_single_test(self, test_definition: TestDefinition) -> Dict[str, Any]:
        """Validate a single test file against its definition"""
        print(f"\n=== VALIDATING: {os.path.basename(test_definition.file_path)} ===")
        print(f"Description: {test_definition.description}")
        
        # ✅ CHECKLIST: Before creating test - Content Definition
        print(f"Expected components: {test_definition.expected_total_count}")
        print(f"Encoding tests: {len(test_definition.encoding_tests)}")
        print(f"Structure tests: {len(test_definition.structure_tests)}")
        
        components = []
        object_id = test_definition.file_path
        
        try:
            # ✅ CHECKLIST: During test execution - Component Analysis
            async for component in self.content_extractor.extract_from_file(test_definition.file_path, object_id):
                components.append(component)
                print(f"  → Component: {component.component_id}")
                print(f"    Type: {component.component_type}")
                print(f"    Method: {getattr(component, 'extraction_method', 'unknown')}")
                print(f"    Content size: {len(component.content) if component.content else 0}")
                if hasattr(component, 'metadata') and component.metadata:
                    if 'original_component_type' in component.metadata:
                        print(f"    Original type: {component.metadata['original_component_type']}")
                print(f"    Content preview: {(component.content or '')[:100]}...")
                print()
        
        except Exception as e:
            print(f"  ERROR during extraction: {e}")
            return {
                "test_definition": test_definition,
                "extraction_error": str(e),
                "validation_results": None
            }
        
        # ✅ CHECKLIST: During test execution - Content Accuracy Verification
        validation_results = self.validator.validate_components(components, test_definition)
        
        # ✅ CHECKLIST: After test completion - Results Validation
        if validation_results["validation_passed"]:
            print(f"  ✅ VALIDATION PASSED")
        else:
            print(f"  ❌ VALIDATION FAILED")
            for error in validation_results["validation_errors"]:
                print(f"     - {error}")
        
        if validation_results["validation_warnings"]:
            print(f"  ⚠️  WARNINGS:")
            for warning in validation_results["validation_warnings"]:
                print(f"     - {warning}")
        
        return {
            "test_definition": test_definition,
            "extraction_error": None,
            "validation_results": validation_results,
            "components_extracted": components
        }
    
    async def run_comprehensive_validation(self):
        """Run complete validation suite"""
        print("=== CONTENT EXTRACTION VALIDATION SUITE ===")
        print("Creating test files with known content and expected outcomes...")
        
        # ✅ CHECKLIST: Before creating tests - Test File Creation Validation
        await self.initialize_extractor()
        
        test_dir = tempfile.mkdtemp(prefix="validation_suite_")
        creator = TestFileCreator(test_dir)
        
        print(f"\nTest directory: {test_dir}")
        
        # Create all test categories
        print("\n1. Creating encoding preservation tests...")
        encoding_tests = creator.create_encoding_test_files()
        print(f"   Created {len(encoding_tests)} encoding tests")
        
        print("\n2. Creating structured format tests...")
        format_tests = creator.create_structured_format_tests()
        print(f"   Created {len(format_tests)} format tests")
        
        print("\n3. Creating office document tests...")
        office_tests = creator.create_office_document_tests()
        print(f"   Created {len(office_tests)} office document tests")
        
        print("\n4. Creating archive tests...")
        archive_tests = creator.create_archive_tests()
        print(f"   Created {len(archive_tests)} archive tests")
        
        print("\n5. Creating OCR tests...")
        ocr_tests = creator.create_ocr_tests()
        print(f"   Created {len(ocr_tests)} OCR tests")
        
        all_test_definitions = creator.get_all_test_definitions()
        print(f"\nTotal tests created: {len(all_test_definitions)}")
        
        # Validate each test
        validation_results = []
        passed_count = 0
        failed_count = 0
        
        for test_def in all_test_definitions:
            result = await self.validate_single_test(test_def)
            validation_results.append(result)
            
            if result["validation_results"] and result["validation_results"]["validation_passed"]:
                passed_count += 1
            else:
                failed_count += 1
        
        # ✅ CHECKLIST: After test completion - Report Accuracy
        self.generate_comprehensive_report(validation_results, passed_count, failed_count)
        
        return validation_results
    
    def generate_comprehensive_report(self, results: List[Dict], passed: int, failed: int):
        """Generate detailed validation report"""
        print("\n" + "="*80)
        print("COMPREHENSIVE VALIDATION REPORT")
        print("="*80)
        
        total_tests = len(results)
        print(f"\nTEST SUMMARY:")
        print(f"  Total tests: {total_tests}")
        print(f"  Passed: {passed} ({passed/total_tests*100:.1f}%)")
        print(f"  Failed: {failed} ({failed/total_tests*100:.1f}%)")
        
        # Category breakdown
        categories = {
            "Encoding Tests": [r for r in results if "utf8" in r["test_definition"].file_path or "utf16" in r["test_definition"].file_path or "win1252" in r["test_definition"].file_path],
            "Format Tests": [r for r in results if "structure" in r["test_definition"].file_path],
            "Office Documents": [r for r in results if "word" in r["test_definition"].file_path or "excel" in r["test_definition"].file_path],
            "Archives": [r for r in results if "archive" in r["test_definition"].file_path],
            "OCR Tests": [r for r in results if "ocr" in r["test_definition"].file_path]
        }
        
        print(f"\nCATEGORY BREAKDOWN:")
        for category, cat_results in categories.items():
            if cat_results:
                cat_passed = sum(1 for r in cat_results if r["validation_results"] and r["validation_results"]["validation_passed"])
                cat_total = len(cat_results)
                print(f"  {category}: {cat_passed}/{cat_total} passed ({cat_passed/cat_total*100:.1f}%)")
        
        # Critical issues analysis
        print(f"\nCRITICAL ISSUES DETECTED:")
        
        encoding_failures = []
        content_duplications = []
        structure_failures = []
        type_misclassifications = []
        
        for result in results:
            if not result["validation_results"]:
                continue
            
            val_result = result["validation_results"]
            
            # Encoding failures
            if val_result["encoding_analysis"]["tests_performed"]:
                for encoding_test in val_result["encoding_analysis"]["tests_performed"]:
                    if not encoding_test.get("encoding_preserved", True):
                        encoding_failures.append({
                            "file": os.path.basename(result["test_definition"].file_path),
                            "test_chars": encoding_test["test_chars"],
                            "encoding": encoding_test["original_encoding"]
                        })
            
            # Content duplications
            if val_result["anti_pattern_analysis"]["detected_issues"]:
                for issue in val_result["anti_pattern_analysis"]["detected_issues"]:
                    if issue["type"] in ["exact_duplicate", "substantial_overlap"]:
                        content_duplications.append({
                            "file": os.path.basename(result["test_definition"].file_path),
                            "issue": issue
                        })
                    elif issue["type"] in ["json_as_text", "csv_as_text"]:
                        type_misclassifications.append({
                            "file": os.path.basename(result["test_definition"].file_path),
                            "issue": issue
                        })
            
            # Structure preservation failures
            if not val_result["validation_passed"]:
                for error in val_result["validation_errors"]:
                    if "structure" in error.lower() or "table" in error.lower() or "json" in error.lower():
                        structure_failures.append({
                            "file": os.path.basename(result["test_definition"].file_path),
                            "error": error
                        })
        
        if encoding_failures:
            print(f"  ❌ ENCODING PRESERVATION FAILURES: {len(encoding_failures)}")
            for failure in encoding_failures[:3]:  # Show first 3
                print(f"    - {failure['file']}: '{failure['test_chars']}' ({failure['encoding']})")
        
        if content_duplications:
            print(f"  ❌ CONTENT DUPLICATIONS DETECTED: {len(content_duplications)}")
            for dup in content_duplications[:3]:
                print(f"    - {dup['file']}: {dup['issue']['type']}")
        
        if type_misclassifications:
            print(f"  ❌ TYPE MISCLASSIFICATIONS: {len(type_misclassifications)}")
            for misclass in type_misclassifications[:3]:
                print(f"    - {misclass['file']}: {misclass['issue']['type']}")
        
        if structure_failures:
            print(f"  ❌ STRUCTURE PRESERVATION FAILURES: {len(structure_failures)}")
            for failure in structure_failures[:3]:
                print(f"    - {failure['file']}: {failure['error'][:80]}...")
        
        if not any([encoding_failures, content_duplications, type_misclassifications, structure_failures]):
            print(f"  ✅ No critical issues detected")
        
        print(f"\n" + "="*80)

async def main():
    """Run the comprehensive validation suite"""
    validator = ContentExtractionValidator()
    results = await validator.run_comprehensive_validation()
    
    # Save detailed results
    output_file = "content_extraction_validation_results.json"
    with open(output_file, "w") as f:
        # Convert results to JSON-serializable format
        serializable_results = []
        for result in results:
            serializable_result = {
                "test_file": result["test_definition"].file_path,
                "description": result["test_definition"].description,
                "extraction_error": result["extraction_error"],
                "validation_results": result["validation_results"]
            }
            serializable_results.append(serializable_result)
        
        json.dump(serializable_results, f, indent=2, default=str)
    
    print(f"\nDetailed validation results saved to: {output_file}")

if __name__ == "__main__":
    asyncio.run(main())